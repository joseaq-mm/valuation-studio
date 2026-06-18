from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SALMON = RGBColor(0x9C, 0x3B, 0x1B)   # acento FT
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Márgenes ajustados para que quepa en ~2 páginas
for s in doc.sections:
    s.top_margin = Inches(0.6)
    s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)

# Estilo base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)
normal.font.color.rgb = DARK


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = SALMON
    p.space_after = Pt(2)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = SALMON
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    return p


def bullet(label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if label:
        r = p.add_run(label + ": ")
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


# ----------------- CONTENIDO -----------------
h1("Valuation Studio — Guía rápida")
sub = doc.add_paragraph()
rs = sub.add_run("Tu Excel de valoración de empresas, convertido en una app que se actualiza sola")
rs.italic = True
rs.font.size = Pt(10.5)
rs.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(6)

body("Valuation Studio nace de un Excel donde introducías a mano datos fundamentales y precios para "
     "obtener ratios que te dicen si una acción está cara o barata. La app hace lo mismo, pero ingiere "
     "los datos automáticamente (Yahoo Finance), aplica TUS fórmulas y los mantiene siempre frescos. "
     "Tiene dos grandes mitades: un módulo CUANTITATIVO (números y ratios) y un módulo CUALITATIVO con IA "
     "(el «Thesis Engine»: tendencias, drivers de crecimiento y tesis de inversión).")

h2("1. ¿Qué puedes hacer? (las pantallas)")
bullet("Empresa", "busca cualquier ticker (AAPL, SAP.DE, SAN.MC…) y ves sus KPIs, ratios clásicos, "
       "métricas de crecimiento, gráficos de ingresos/FCF y los precios objetivo POC/POV. Puedes editar "
       "los inputs a mano y recalcular en vivo.")
bullet("Watchlist", "lista de seguimiento con snapshots manuales (guardas tus propios ajustes) y alertas "
       "por fila (campana) cuando una empresa cruza de cara a barata.")
bullet("Cartera", "tus posiciones reales con P/L vivo, P/L % y KPIs totales; también posiciones de «solo seguimiento».")
bullet("Comparar", "hasta 6 empresas lado a lado.")
bullet("Tesis (Thesis Engine)", "el cerebro cualitativo: explora tendencias, genera tesis de inversión por "
       "empresa y agrupa todo en «Megatesis». Más detalle abajo.")

h2("2. Los ratios y por qué existen")
body("Todo gira en torno a dos precios objetivo calculados con TUS fórmulas:")
bullet("POC — Precio Objetivo de Compra", "estima lo que «vale» la acción hoy combinando ingresos por "
       "acción, margen bruto, caja libre frente a deuda y capitalización, y el crecimiento esperado de "
       "ingresos y FCF. Es tu ancla: si el precio actual está muy por debajo del POC, hay margen de compra.")
bullet("Ratio de Compra (%)", "cuánto separa el precio actual del POC. Positivo y alto = potencialmente barata.")
bullet("POV — Precio Objetivo de Venta", "el POC ajustado por el margen operativo; marca la zona donde la "
       "acción empieza a estar exigente.")
bullet("Ratio de Venta (%)", "distancia del precio actual al POV; te avisa de cuándo el recorrido se agota.")
body("Además verás ratios clásicos (P/E, PEG, P/B, P/S, EV/EBITDA, ROE, ROA, márgenes, deuda/equity…) y una "
     "sección de crecimiento (crecimiento de ingresos, CAGR, margen de FCF, Rule of 40 y año estimado de "
     "breakeven). Cada fila tiene un punto de salud verde/ámbar/rojo y un tooltip que explica qué mide y sus "
     "rangos «barato/normal/caro». Las proyecciones se generan solas, pero se avisa (en ámbar) cuando un "
     "crecimiento es tan extremo que conviene revisarlo a mano.")

h2("3. El Thesis Engine: tendencias, drivers y tesis")
body("Aquí la IA (vía Emergent LLM) busca información en vivo en la web y estructura el análisis. Conceptos clave:")
bullet("Tendencia", "un tema o megatendencia (IA, vehículos eléctricos, energía nuclear…). Al explorarla ves "
       "su cadena de valor con LÍDERES, COMPETIDORES y DISRUPTORES, además de su TAM (tamaño de mercado) y CAGR.")
bullet("Tesis", "una apuesta de alta convicción sobre un DRIVER DE CRECIMIENTO independiente de una empresa. "
       "Cada driver tiene su propio TAM que NO se solapa con otros, para no contar dos veces el mismo mercado. "
       "Se etiquetan como «Crecimiento actual» o «Apuesta futura».")
bullet("Planificar → Ejecutar", "primero organizas los drivers (fusionar parecidos, partir uno grande en "
       "sub-partes) hasta dejar el TAM 100% mutuamente excluyente; luego generas todas las tesis de golpe.")
bullet("Tendencia / Tesis automática", "botones que descubren temas emergentes por momentum para que no "
       "partas de cero.")

h2("4. Los scores del Thesis Engine y su porqué")
bullet("Score global tendencia (0–100)", "calidad cualitativa de la empresa dentro de una tendencia. Combina "
       "cuatro sub-scores: posición competitiva, momentum del sector, calidad del management y resiliencia "
       "financiera. Responde a «¿qué tan buena es esta empresa en este tema?».")
bullet("Probabilidad (0–10)", "certeza de que la tendencia realmente ocurra, según la evidencia de las fuentes. "
       "Verde = muy probable. La «contratesis» mide la consecuencia derivada (quién pierde si la tendencia avanza).")
bullet("TAM Score", "el puente entre lo cualitativo y lo cuantitativo. Mezcla la calidad (score global) con el "
       "trozo de mercado (TAM del eslabón) que le toca a la empresa frente a sus ingresos proyectados. Mayor que "
       "1 (verde) = la oportunidad es grande respecto a su tamaño actual; menor que 1 (ámbar) = más modesta. "
       "Sirve para priorizar de un vistazo dónde hay más recorrido.")
bullet("Win probability (0–10)", "probabilidad de que la empresa sea ganadora en ese driver concreto; ayuda a "
       "ordenar las apuestas con más momentum.")

h2("5. Cómo empezar en 5 minutos")
bullet("1", "Busca una empresa que conozcas y mira su POC, POV y los puntos de salud de los ratios.")
bullet("2", "Edita algún input (p. ej. el crecimiento) y pulsa recalcular para ver cómo se mueve el POC.")
bullet("3", "Añádela a tu Watchlist y activa la campana para recibir alertas cuando cruce a «barata».")
bullet("4", "Ve a «Tesis», explora una tendencia o genera una tesis desde una empresa y observa los scores.")
bullet("5", "Inicia sesión con Google para sincronizar watchlist, cartera y tesis entre dispositivos, y activa "
       "el Screener nocturno / Radar semanal por email.")

# Nota de cierre
note = doc.add_paragraph()
rn = note.add_run("Idea clave: los números (POC/POV/ratios) te dicen si está BARATA; el Thesis Engine te dice "
                  "si el negocio va a CRECER. Las mejores oportunidades aciertan en las dos cosas a la vez.")
rn.italic = True
rn.font.size = Pt(9.5)
rn.font.color.rgb = GREY
note.paragraph_format.space_before = Pt(8)

doc.save('/app/Valuation_Studio_Guia_Rapida.docx')
print('OK guardado')
