"""Builds the comprehensive Valuation Studio user manual (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xB3, 0x2A, 0x22)
NAVY = RGBColor(0x05, 0x20, 0x49)
GREY = RGBColor(0x4A, 0x4A, 0x4A)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = RED
    r.font.size = Pt(18)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(14)
    return p


def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(12)
    return p


def para(text, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(label, text=None):
    p = doc.add_paragraph(style="List Bullet")
    if text:
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(label)
    return p


def numbered(text):
    doc.add_paragraph(text, style="List Number")


# ---------------- Portada ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Valuation Studio")
tr.font.size = Pt(30)
tr.font.color.rgb = RED
tr.bold = True
sub = doc.add_paragraph()
sr = sub.add_run("Manual de uso — todas las páginas, funcionalidades y filosofía")
sr.italic = True
sr.font.size = Pt(13)
sr.font.color.rgb = GREY

para(
    "Valuation Studio nace de un Excel de valoración fundamental de empresas. En ese Excel "
    "introducías a mano datos y precios para obtener ratios que te decían si una acción estaba "
    "cara o barata. La app hace lo mismo pero ingiere los datos automáticamente (Yahoo Finance), "
    "aplica TUS fórmulas y los mantiene siempre frescos. Además añade un cerebro cualitativo con "
    "IA (el «Thesis Engine») y un módulo macro que mide si el mercado, en su conjunto, está caro o barato."
)

# ---------------- Filosofía ----------------
h1("1. La filosofía: dos preguntas, una decisión")
para(
    "Toda la app gira en torno a responder dos preguntas independientes sobre cada empresa, y a "
    "juntarlas para decidir:"
)
bullet("¿Está BARATA?", "es la parte CUANTITATIVA. Se responde con precios objetivo y ratios calculados con tus fórmulas (POC, POV, Ratio de Compra/Venta).")
bullet("¿Va a CRECER y es BUEN NEGOCIO?", "es la parte CUALITATIVA (con IA). Se responde con tendencias, tesis de crecimiento, scores de calidad y la validación operativa vía KPIs.")
para(
    "Idea clave: los números te dicen si está barata; el análisis cualitativo te dice si el negocio "
    "va a crecer. Las mejores oportunidades aciertan en las DOS cosas a la vez. Encima de todo, el "
    "módulo MACRO te dice si conviene ser más agresivo (mercado barato → crecimiento) o más "
    "defensivo (mercado caro → efectivo/defensivas), para modular cuánto peso das a cada tipo de empresa.",
    italic=True,
)
para(
    "Filosofía de niveles: organizas tus empresas en Nivel 1 (tus preferidas y posiciones reales) "
    "y Nivel 2 (interesantes, en seguimiento). Sobre ellas construyes tesis y KPIs, y las vigilas en "
    "el tiempo, reescribiendo lo que envejece o cuando hay noticias/resultados relevantes."
)

# ---------------- Las páginas ----------------
h1("2. Las páginas y sus funcionalidades")

h2("Inicio / Buscador")
para("Punto de entrada. Busca cualquier ticker (AAPL, SAP.DE, SAN.MC…) desde la barra superior, que está disponible en todas las páginas. Al iniciar sesión con Google, todo se sincroniza entre dispositivos y se activan las alertas por email.")

h2("Empresa (ficha individual)")
para("La radiografía cuantitativa de una acción. Contiene:")
bullet("POC y POV", "los dos precios objetivo (compra y venta) calculados con tus fórmulas; ver la sección de conceptos.")
bullet("Ratios clásicos", "P/E, PEG, P/B, P/S, EV/EBITDA, ROE, ROA, márgenes, deuda/equity… cada fila con un punto de salud verde/ámbar/rojo y un tooltip que explica qué mide y sus rangos barato/normal/caro.")
bullet("Crecimiento", "CAGR de ingresos y FCF, margen de FCF, Rule of 40, breakeven. Las proyecciones se generan solas y avisan en ámbar cuando un crecimiento es extremo y conviene revisarlo a mano.")
bullet("Edición en vivo", "puedes editar los inputs (p. ej. el crecimiento esperado) y recalcular al instante para ver cómo se mueven POC/POV.")
bullet("Añadir a Nivel 1 / Nivel 2", "desde la ficha incorporas la empresa a tus listas.")

h2("Nivel 1 (tus preferidas + cartera real)")
para("Tus acciones de mayor convicción y tus posiciones reales, con P/L vivo, P/L % y KPIs totales. Admite también posiciones de «solo seguimiento» (sin dinero invertido).")

h2("Nivel 2 (seguimiento)")
para("Acciones que te interesan pero no son tus preferidas. Lista de seguimiento con snapshots manuales y alertas por fila (campana) que te avisan cuando una empresa cruza de cara a barata.")

h2("Comparar")
para("Hasta 6 empresas lado a lado para contrastar ratios, crecimiento y precios objetivo de un vistazo.")

h2("Tesis (Thesis Engine) — el cerebro cualitativo")
para("Aquí la IA busca información en vivo y estructura el análisis. Tiene dos modos de trabajo:")
bullet("Modo Empresa", "generas la tesis de una empresa concreta: sus drivers de crecimiento, scores y valoración cualitativa.")
bullet("Modo Explorar (Tendencias → Empresas)", "partes de una tendencia/tema y descubres las empresas de su cadena de valor (líderes, competidores, disruptores), su TAM y CAGR, y hasta ETFs/fondos de la temática.")
para("Otras funciones importantes:")
bullet("Tendencia automática", "con un clic descubre un tema emergente por momentum, para no empezar de cero.")
bullet("Planificar → Ejecutar", "primero organizas los drivers (fusionar parecidos, partir uno grande) hasta que el TAM quede 100% mutuamente excluyente (sin contar dos veces el mismo mercado); luego generas todas las tesis de golpe.")
bullet("Megatesis", "agrupas varias tesis/tendencias relacionadas en una carpeta (megatendencia).")
bullet("Deshacer/Rehacer y Refrescar datos", "puedes revertir cambios y volver a leer los fundamentales de Yahoo para recalcular los TAM Score con números frescos (no toca los scores cualitativos).")

h3("El mapa (treemap) del explorador de Tesis")
para("Dentro de Tesis, el explorador muestra tus datos como un mapa de rectángulos (treemap): cuanto más grande y más verde el rectángulo, mejor en la métrica elegida. Tiene 5 vistas:")
bullet("Megatendencias", "tus carpetas. Tamaño = media del CAGR a 4 años de sus tendencias; muestra el TAM total (suma).")
bullet("Tendencias", "tus tendencias. Con un desplegable eliges qué las dimensiona: por CAGR, por TAM, o por la media normalizada de ambas (para ver cuáles destacan en las dos variables a la vez).")
bullet("Convergencia", "empresas que aparecen en VARIAS de tus tendencias a la vez (señal fuerte). El tamaño es el nº de tendencias; las que están en 2 o más llevan borde dorado.")
bullet("Empresas · score medio", "tus empresas totalmente desarrolladas, por su score global medio.")
bullet("Empresas · TAM total", "las mismas, por la suma de sus TAM Score.")

h2("Visual — mapa de oportunidades")
para("Cruza en una sola pantalla lo cuantitativo y lo cualitativo de todas las empresas que son miembro de tus tesis. Dos elementos:")
bullet("El gráfico (scatter/cuadrante)", "cada burbuja es una empresa. Ver la sección «Cómo leer el gráfico de Visual».")
bullet("La tabla", "todas las empresas con Score, TAM Score, Coef KPI, Combinado cualitativo, Ratio Compra %, Ratio Venta % y Combinado total. Ordenable por cualquier columna.")
para("Herramientas de Visual:")
bullet("Filtros", "Score, TAM Score, Coef KPI, Ratio Compra, Ratio Venta, Combinado cualitativo y Combinado total. Al aplicarlos, además de filtrar el gráfico, marcan/desmarcan automáticamente las filas de la tabla.")
bullet("Campana de alerta (por empresa)", "junto al nombre. Al pulsarla configuras umbrales de Score, TAM Score y Coef KPI (con dirección ≥ o ≤). Si la empresa alcanza esa situación, recibes un email ese día. Las empresas con campana también avisan al cruzar de barato↔caro. Todo se consolida en UN único email diario para no saturarte.")

h2("KPIs — validación operativa")
para("Comprueba si lo que promete la tesis se está cumpliendo en los números reales de la empresa. Funciones:")
bullet("Extracción de documentos", "sube o deja que la app busque informes/PDF de la empresa y extrae los KPIs operativos automáticamente (con reintentos y lectura local de PDF).")
bullet("Coeficiente KPI", "un número que resume si los KPIs se están validando (por encima de su punto neutro) o deteriorando (por debajo).")
bullet("Fuentes trazables", "cada KPI enlaza a su documento de origen y muestra la cita textual en un tooltip.")
bullet("Indicador de tendencia", "junto al coeficiente, una flecha ↑/↓ indica cómo ha cambiado respecto al análisis anterior (histórico de reanálisis).")

h2("Macro — ¿el mercado está caro o barato?")
para("Calcula un coeficiente de mercado en tiempo real para saber si conviene ser agresivo o defensivo. Contiene:")
bullet("Indicadores", "Renta variable, PIB, M3 proxy, Tipo FED, Inflación, Productividad, Petróleo y Mix energético. Renta variable y PIB muestran el valor oficial y, en grande, una estimación EN VIVO (la renta variable se ajusta por el índice elegido: S&P 500, NASDAQ 100 o Dow Jones; el PIB por su crecimiento interanual prorrateado).")
bullet("Coeficiente de mercado (aguja/dial)", "por debajo de 1 = mercado caro (sesgo defensivo/efectivo); por encima de 1 = barato (sesgo crecimiento). Al lado se muestra la diferencia hasta 1 en %, y dentro del cuadro caro/barato correspondiente.")
bullet("Petróleo con deslizador", "compara el precio actual con su media histórica (eliges los años).")
bullet("Mix energético mundial", "peso de cada fuente de energía.")
bullet("Gráfico de evolución (10 años)", "Renta variable, PIB, su resta y Productividad en el tiempo, con doble eje Y; el último punto es estimado. Ampliable.")
bullet("Histórico del coeficiente", "mini-gráfico con zonas sombreadas (verde = barato, rojo = caro) y línea neutra en 1. Ampliable.")

h2("Instrucciones")
para("Esta misma guía dentro de la app, con opción de descarga.")

# ---------------- Conceptos ----------------
h1("3. Conceptos clave (glosario)")
bullet("POC — Precio Objetivo de Compra", "lo que «vale» la acción hoy según tus fórmulas (ingresos por acción, margen bruto, caja libre frente a deuda y capitalización, y crecimiento esperado). Tu ancla de valor.")
bullet("Ratio de Compra (%)", "distancia del precio actual al POC. Positivo y alto = potencialmente barata.")
bullet("POV — Precio Objetivo de Venta", "el POC ajustado por el margen operativo; zona donde la acción empieza a estar exigente.")
bullet("Ratio de Venta (%)", "distancia del precio actual al POV; avisa de cuándo se agota el recorrido.")
bullet("Score global de tendencia (0–100)", "calidad cualitativa de la empresa en un tema. Combina posición competitiva, momentum del sector, calidad del management y resiliencia financiera.")
bullet("TAM Score", "el puente entre lo cualitativo y lo cuantitativo: mezcla la calidad con el trozo de mercado (TAM) que le toca a la empresa frente a sus ingresos proyectados. >1 (verde) = oportunidad grande respecto a su tamaño; <1 (ámbar) = más modesta.")
bullet("Coeficiente KPI", "resume si los KPIs operativos validan (>punto neutro) o deterioran (<) la tesis.")
bullet("Combinado cualitativo", "une el Score con el Coef KPI (calidad + validación operativa).")
bullet("Combinado total", "une lo cualitativo con lo cuantitativo (precio). Es la nota global de oportunidad.")
bullet("Tendencia", "un tema/megatendencia (IA, vehículos eléctricos, energía nuclear…), con su cadena de valor, TAM y CAGR.")
bullet("Tesis / Driver de crecimiento", "una apuesta de alta convicción sobre un motor de crecimiento con su propio TAM que no se solapa con otros.")
bullet("TAM y CAGR", "TAM = tamaño del mercado total abordable; CAGR = crecimiento compuesto anual estimado.")
bullet("Convergencia", "una empresa que aparece en varias de tus tendencias a la vez: señal de que está en el cruce de varias megatendencias.")
bullet("Coeficiente de mercado (Macro)", "medida global cara/barata del mercado que orienta tu sesgo agresivo/defensivo.")

# ---------------- Relaciones ----------------
h1("4. Cómo se relacionan las páginas")
para("El flujo natural conecta todo:")
numbered("Descubres ideas en Tesis (explorando tendencias o con tendencia automática) o buscas empresas concretas en Empresa/Tesis.")
numbered("Clasificas las que te convencen en Nivel 1 y Nivel 2.")
numbered("Desarrollas su tesis (cualitativo) y sus KPIs (validación), que alimentan el Score, el TAM Score y el Coef KPI.")
numbered("Todas esas empresas y métricas aparecen juntas en Visual (gráfico + tabla), donde comparas y filtras.")
numbered("Macro te dice el clima general (caro/barato) para modular cuánto peso dar a cada tipo de empresa y cuánto efectivo mantener.")
numbered("Pones campanas de alerta en Visual y recibes un email diario cuando algo cambia. Reescribes tesis/KPIs cuando envejecen o hay noticias/resultados.")

# ---------------- Cómo leer el gráfico de Visual ----------------
h1("5. Cómo leer el gráfico de Visual")
para("El gráfico de Visual es un cuadrante tipo BCG. Cada burbuja es una empresa:")
bullet("Eje horizontal (X) = Score cualitativo", "a la derecha, mejor calidad de negocio.")
bullet("Eje vertical (Y) = Ratio de Compra %", "más arriba, más barata (más margen hasta el POC).")
bullet("Tamaño de la burbuja = TAM Score", "más grande, mayor oportunidad de mercado relativa.")
bullet("Color", "refleja la señal barato/normal/caro según tus umbrales.")
bullet("Líneas discontinuas (medianas)", "dividen el gráfico en 4 cuadrantes respecto a la mediana de Score y de Ratio de Compra.")
para(
    "Lectura práctica: el cuadrante SUPERIOR DERECHO (alta calidad + barata) es el de las mejores "
    "oportunidades (buen negocio a buen precio). El inferior izquierdo (baja calidad + cara) es el "
    "menos atractivo. Arriba-izquierda = barata pero de menor calidad (cautela); abajo-derecha = "
    "gran negocio pero caro (paciencia / lista de espera).",
    italic=True,
)
para(
    "Diferencia con el mapa (treemap) de Tesis: el treemap sirve para ver el PESO relativo de tus "
    "tendencias/empresas (tamaño = magnitud de una métrica; verde→rojo = alto→bajo). El gráfico de "
    "Visual sirve para POSICIONAR cada empresa en calidad vs. precio. Uno mide magnitud; el otro, posición."
)

# ---------------- Ejemplos ----------------
h1("6. Dos formas de usar la app")

h2("A) Inversor que ya tiene cartera")
numbered("Busca sus acciones y las clasifica en listas de Nivel 1 (preferidas/posiciones) y Nivel 2 (seguimiento).")
numbered("Desarrolla la tesis de cada empresa y sus KPIs, obteniendo Score, TAM Score y Coef KPI.")
numbered("Consulta Macro para decidir el peso de cada tipología de empresa y el nivel de efectivo (más defensivo si el mercado está caro; más crecimiento si está barato).")
numbered("Vigila la evolución: reescribe las tesis y KPIs más antiguos, o cuando hay noticias/resultados relevantes.")
numbered("Está atento a nuevas oportunidades y busca ideas con «Tendencias → Empresas», donde también aparecen ETFs de cada temática.")
numbered("Pone campanas de alerta en Visual sobre sus empresas clave y recibe el email diario consolidado.")

h2("B) Inversor que empieza de cero")
numbered("Explora primero con las tendencias (o con «Tendencia automática») para descubrir temas y empresas, incluidos ETFs.")
numbered("Alternativamente, escribe empresas directamente en el buscador para ver su Ratio de Compra/Venta (precios) y su tesis (idea cualitativa).")
numbered("Con esa información monta sus propias listas de Nivel 1 y/o Nivel 2 para decidir dónde invertir.")
numbered("A partir de ahí sigue el mismo camino que el inversor experimentado: desarrolla tesis y KPIs, consulta Macro, vigila con alertas y actualiza lo que envejece.")

para(
    "En ambos casos, el objetivo es el mismo: encontrar buenos negocios (cualitativo) a buen precio "
    "(cuantitativo), en el momento adecuado del mercado (macro), y no perderles la pista (alertas).",
    italic=True,
    color=GREY,
)

doc.save("/app/frontend/public/Valuation_Studio_Manual.docx")
print("saved")
