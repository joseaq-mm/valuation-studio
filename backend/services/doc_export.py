"""Server-side export of theses (company / tendencia) and KPI text docs to PDF and
Word (.docx). We build a neutral list of "blocks" from the source document and render
it to both formats, so the content logic lives in one place."""
import io
import re
from html import escape

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = colors.HexColor("#052049")
GREY = colors.HexColor("#4A4A4A")


# ── Formatting helpers ────────────────────────────────────────────────────────
def _busd(v):
    try:
        v = float(v)
    except (TypeError, ValueError):
        return "—"
    if v <= 0:
        return "—"
    if v >= 1000:
        return f"${v/1000:.1f} T"
    return f"${v:.0f} B"


def _txt(v):
    return "" if v is None else str(v).strip()


def safe_filename(name, ext):
    base = re.sub(r"[^\w\-. áéíóúñÁÉÍÓÚÑ]", "", (name or "documento")).strip()
    base = re.sub(r"\s+", " ", base)[:80] or "documento"
    return f"{base}.{ext}"


# ── Blocks (neutral intermediate representation) ──────────────────────────────
def thesis_blocks(doc):
    t = doc.get("type")
    return _company_blocks(doc) if t == "company" else _tendencia_blocks(doc)


def _company_blocks(doc):
    comp = doc.get("company") or {}
    ticker = _txt(comp.get("ticker"))
    blocks = []
    blocks.append({"h1": _txt(doc.get("title")) or _txt(comp.get("name")) or "Tesis"})
    meta = []
    if ticker:
        meta.append(("Ticker", ticker))
    if doc.get("probability") is not None:
        meta.append(("Probabilidad de tesis", f"{doc.get('probability')}/10"))
    if doc.get("overall_relevance") is not None:
        meta.append(("Relevancia global", f"{doc.get('overall_relevance')}/100"))
    meta.append(("Tipo", "Ficha de empresa (Nivel 2)"))
    if meta:
        blocks.append({"kv": meta})
    if _txt(doc.get("summary")):
        blocks.append({"h2": "Resumen"})
        blocks.append({"p": _txt(doc.get("summary"))})
    if _txt(doc.get("probability_rationale")):
        blocks.append({"h2": "Justificación de la probabilidad"})
        blocks.append({"p": _txt(doc.get("probability_rationale"))})

    trends = doc.get("trends") or []
    if trends:
        blocks.append({"h2": "Drivers de crecimiento / tesis"})
        rows = [["Driver", "Tipo", "TAM", "Relevancia", "Prob. victoria"]]
        for tr in trends:
            rows.append([
                _txt(tr.get("name")),
                _txt(tr.get("type")),
                _busd(tr.get("tam_busd")),
                f"{tr.get('relevance_score')}" if tr.get("relevance_score") is not None else "—",
                f"{tr.get('win_probability')}%" if tr.get("win_probability") is not None else "—",
            ])
        blocks.append({"table": {"headers": rows[0], "rows": rows[1:]}})
        for tr in trends:
            blocks.append({"h3": _txt(tr.get("name"))})
            if _txt(tr.get("demand_driver")):
                blocks.append({"p": f"Driver de demanda: {_txt(tr.get('demand_driver'))}"})
            if _txt(tr.get("fit_description")):
                blocks.append({"p": f"Encaje: {_txt(tr.get('fit_description'))}"})
            if _txt(tr.get("rationale")):
                blocks.append({"p": _txt(tr.get("rationale"))})
    _append_sources(blocks, doc.get("sources"))
    return blocks


def _tendencia_blocks(doc):
    blocks = []
    blocks.append({"h1": _txt(doc.get("title")) or "Tendencia"})
    tam = doc.get("tam") or {}
    meta = [("Tipo", "Tendencia (Nivel 2)")]
    if tam.get("global_busd") is not None:
        yr = f" ({tam.get('year')})" if tam.get("year") else ""
        meta.append(("TAM global", f"{_busd(tam.get('global_busd'))}{yr}"))
    if doc.get("cagr_4y") is not None:
        meta.append(("CAGR 4 años", f"{doc.get('cagr_4y')}%"))
    blocks.append({"kv": meta})
    if _txt(doc.get("summary")):
        blocks.append({"h2": "Resumen"})
        blocks.append({"p": _txt(doc.get("summary"))})
    if _txt(doc.get("cagr_note")):
        blocks.append({"h2": "Nota sobre el crecimiento"})
        blocks.append({"p": _txt(doc.get("cagr_note"))})

    vc = doc.get("value_chain") or []
    if vc:
        blocks.append({"h2": "Cadena de valor"})
        rows = [["Etapa", "TAM", "Descripción"]]
        for s in vc:
            rows.append([_txt(s.get("stage")), _busd(s.get("tam_busd")), _txt(s.get("description"))])
        blocks.append({"table": {"headers": rows[0], "rows": rows[1:], "wrap_last": True}})

    comps = doc.get("companies") or []
    if comps:
        blocks.append({"h2": "Empresas de la tendencia"})
        rows = [["Empresa", "Ticker", "Rol", "Categoría"]]
        for c in comps:
            rows.append([_txt(c.get("name")), _txt(c.get("ticker")), _txt(c.get("value_chain_role")), _txt(c.get("category"))])
        blocks.append({"table": {"headers": rows[0], "rows": rows[1:]}})
        for c in comps:
            if _txt(c.get("why")):
                blocks.append({"p": f"{_txt(c.get('name'))} ({_txt(c.get('ticker'))}): {_txt(c.get('why'))}"})
    _append_sources(blocks, doc.get("sources"))
    return blocks


def _append_sources(blocks, sources):
    sources = sources or []
    if not sources:
        return
    blocks.append({"h2": "Fuentes"})
    items = []
    for s in sources[:25]:
        title = _txt(s.get("title")) or _txt(s.get("url"))
        url = _txt(s.get("url"))
        items.append(f"{title} — {url}" if url else title)
    blocks.append({"bullets": items})


def kpi_doc_blocks(title, text):
    """Blocks for a saved KPI chat/text document."""
    blocks = [{"h1": _txt(title) or "Documento KPI"}]
    for para in re.split(r"\n{2,}", _txt(text)):
        p = para.strip()
        if p:
            blocks.append({"p": p})
    return blocks


def _coef_label(c):
    try:
        c = float(c)
    except (TypeError, ValueError):
        return "Sin datos"
    if c >= 1.05:
        return "Validándose"
    if c <= 0.95:
        return "Deteriorándose"
    return "Neutral"


def _num(v, dec=2):
    try:
        return f"{float(v):.{dec}f}"
    except (TypeError, ValueError):
        return "—"


def kpi_snapshot_blocks(company_name, snap):
    """Blocks for a company's operational-KPI validation snapshot (coefficient,
    per-driver verdicts, the KPI table and the sources)."""
    snap = snap or {}
    name = _txt(company_name) or "Empresa"
    blocks = [{"h1": f"Validación de KPIs · {name}"}]

    coef = snap.get("coef_global")
    meta = [("Coeficiente global", f"{_num(coef)} ({_coef_label(coef)})")]
    if snap.get("signal_global") is not None:
        meta.append(("Señal agregada (S)", _num(snap.get("signal_global"))))
    if snap.get("period"):
        meta.append(("Periodo", _txt(snap.get("period"))))
    if snap.get("generated_at"):
        meta.append(("Analizado", _txt(snap.get("generated_at"))[:10]))
    meta.append(("Tipo", "Validación operativa (Nivel 2)"))
    blocks.append({"kv": meta})

    if _txt(snap.get("note")):
        blocks.append({"p": _txt(snap.get("note"))})

    drivers = snap.get("drivers") or []
    if drivers:
        blocks.append({"h2": "Coeficiente por driver"})
        rows = [["Driver", "Coeficiente", "Nº KPIs", "Veredicto"]]
        for d in drivers:
            rows.append([
                _txt(d.get("name")),
                f"{_num(d.get('coef'))} ({_coef_label(d.get('coef'))})",
                str(d.get("n_kpis") if d.get("n_kpis") is not None else "—"),
                _txt(d.get("verdict")),
            ])
        blocks.append({"table": {"headers": rows[0], "rows": rows[1:], "wrap_last": True}})

    kpis = snap.get("kpis") or []
    if kpis:
        blocks.append({"h2": "KPIs operativos"})
        rows = [["KPI", "Actual", "Anterior", "YoY", "Señal", "Peso", "Driver"]]
        for k in kpis:
            rows.append([
                _txt(k.get("name")),
                _txt(k.get("value_current")) or "—",
                _txt(k.get("value_prior")) or "—",
                _txt(k.get("yoy_change")) or "—",
                _num(k.get("signal")) if k.get("signal") is not None else "—",
                _num(k.get("weight")) if k.get("weight") is not None else "—",
                _txt(k.get("driver")),
            ])
        blocks.append({"table": {"headers": rows[0], "rows": rows[1:]}})
        for k in kpis:
            if _txt(k.get("rationale")):
                blocks.append({"p": f"{_txt(k.get('name'))}: {_txt(k.get('rationale'))}"})

    _append_sources(blocks, snap.get("sources"))
    return blocks


def kpi_snapshot_pdf(company_name, snap):
    return render_pdf(kpi_snapshot_blocks(company_name, snap), f"KPIs {company_name}")


def kpi_snapshot_docx(company_name, snap):
    return render_docx(kpi_snapshot_blocks(company_name, snap), f"KPIs {company_name}")


# ── Renderers ─────────────────────────────────────────────────────────────────
def render_pdf(blocks, doc_title=""):
    buf = io.BytesIO()
    docpdf = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=18 * mm,
                               leftMargin=18 * mm, rightMargin=18 * mm, title=doc_title or "Documento")
    ss = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=ss["Title"], textColor=NAVY, fontSize=20, spaceAfter=6, alignment=0)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], textColor=NAVY, fontSize=13, spaceBefore=12, spaceAfter=4)
    h3 = ParagraphStyle("h3", parent=ss["Heading3"], textColor=GREY, fontSize=11, spaceBefore=8, spaceAfter=2)
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=13, spaceAfter=4)
    cell = ParagraphStyle("cell", parent=body, fontSize=8.5, leading=11, spaceAfter=0)

    def esc(x):
        return escape(str(x)).replace("\n", "<br/>")

    story = []
    for b in blocks:
        if "h1" in b:
            story.append(Paragraph(esc(b["h1"]), h1))
            story.append(Spacer(1, 4))
        elif "h2" in b:
            story.append(Paragraph(esc(b["h2"]), h2))
        elif "h3" in b:
            story.append(Paragraph(esc(b["h3"]), h3))
        elif "p" in b:
            story.append(Paragraph(esc(b["p"]), body))
        elif "kv" in b:
            data = [[Paragraph(f"<b>{esc(k)}</b>", cell), Paragraph(esc(v), cell)] for k, v in b["kv"]]
            tbl = Table(data, colWidths=[45 * mm, None])
            tbl.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#E0E0E0")),
                ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 6))
        elif "table" in b:
            t = b["table"]
            data = [[Paragraph(f"<b>{esc(h)}</b>", cell) for h in t["headers"]]]
            for r in t["rows"]:
                data.append([Paragraph(esc(c), cell) for c in r])
            tbl = Table(data, repeatRows=1, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F6FA")]),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 8))
        elif "bullets" in b:
            items = [ListItem(Paragraph(esc(x), body), leftIndent=8) for x in b["bullets"]]
            story.append(ListFlowable(items, bulletType="bullet", start="•"))
            story.append(Spacer(1, 6))

    docpdf.build(story)
    return buf.getvalue()


def render_docx(blocks, doc_title=""):
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for b in blocks:
        if "h1" in b:
            h = d.add_heading(b["h1"], level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x05, 0x20, 0x49)
        elif "h2" in b:
            h = d.add_heading(b["h2"], level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x05, 0x20, 0x49)
        elif "h3" in b:
            d.add_heading(b["h3"], level=2)
        elif "p" in b:
            d.add_paragraph(b["p"])
        elif "kv" in b:
            tbl = d.add_table(rows=0, cols=2)
            tbl.style = "Light Grid Accent 1"
            for k, v in b["kv"]:
                cells = tbl.add_row().cells
                cells[0].text = str(k)
                cells[1].text = str(v)
                for r in cells[0].paragraphs[0].runs:
                    r.bold = True
        elif "table" in b:
            t = b["table"]
            tbl = d.add_table(rows=1, cols=len(t["headers"]))
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            for i, h in enumerate(t["headers"]):
                hdr[i].text = str(h)
                for r in hdr[i].paragraphs[0].runs:
                    r.bold = True
            for row in t["rows"]:
                cells = tbl.add_row().cells
                for i, c in enumerate(row):
                    cells[i].text = str(c)
        elif "bullets" in b:
            for x in b["bullets"]:
                d.add_paragraph(str(x), style="List Bullet")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def thesis_pdf(doc):
    return render_pdf(thesis_blocks(doc), _txt(doc.get("title")))


def thesis_docx(doc):
    return render_docx(thesis_blocks(doc), _txt(doc.get("title")))


# ── KPI source-document conversion (original → chosen output format) ───────────
# The uploaded document (PDF/image) is the source of truth; text docs (transcripts,
# auto-fetched web/SEC) use their stored text. We convert the ORIGINAL to the format
# the user picks (PDF · Word · JPG). JPG is only offered for image originals.

def image_to_jpg(data):
    """Return the image as JPEG bytes (flattening transparency onto white)."""
    from PIL import Image
    im = Image.open(io.BytesIO(data))
    if im.mode in ("RGBA", "LA", "P"):
        bg = Image.new("RGB", im.size, (255, 255, 255))
        im = im.convert("RGBA")
        bg.paste(im, mask=im.split()[-1])
        im = bg
    else:
        im = im.convert("RGB")
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=92)
    return buf.getvalue()


def image_to_pdf(data):
    """Embed an image on an A4 page, scaled to fit with margins."""
    from PIL import Image
    from reportlab.pdfgen import canvas as _canvas
    im = Image.open(io.BytesIO(data)).convert("RGB")
    buf = io.BytesIO()
    c = _canvas.Canvas(buf, pagesize=A4)
    pw, ph = A4
    margin = 15 * mm
    max_w, max_h = pw - 2 * margin, ph - 2 * margin
    iw, ih = im.size
    ratio = min(max_w / iw, max_h / ih)
    dw, dh = iw * ratio, ih * ratio
    x, y = (pw - dw) / 2, (ph - dh) / 2
    from reportlab.lib.utils import ImageReader
    c.drawImage(ImageReader(im), x, y, width=dw, height=dh, preserveAspectRatio=True, mask="auto")
    c.showPage()
    c.save()
    return buf.getvalue()


def image_to_docx(data, title=None):
    """Embed an image in a .docx, scaled to page width."""
    jpg = image_to_jpg(data)
    d = Document()
    if title:
        d.add_heading(_txt(title)[:120] or "Documento", level=1)
    d.add_picture(io.BytesIO(jpg), width=Inches(6.2))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def pdf_to_docx(data, title=None):
    """Faithful copy of a PDF into a .docx by rendering each page to an image
    (up to 40 pages) and embedding it full width."""
    import fitz  # PyMuPDF
    d = Document()
    if title:
        d.add_heading(_txt(title)[:120] or "Documento", level=1)
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        zoom = 150 / 72.0  # ~150 DPI
        mat = fitz.Matrix(zoom, zoom)
        for i, page in enumerate(doc):
            if i >= 40:
                break
            pix = page.get_pixmap(matrix=mat, alpha=False)
            d.add_picture(io.BytesIO(pix.tobytes("png")), width=Inches(6.4))
    finally:
        doc.close()
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def text_to_pdf(title, text):
    return render_pdf(kpi_doc_blocks(title, text), _txt(title) or "Documento")


def text_to_docx(title, text):
    return render_docx(kpi_doc_blocks(title, text), _txt(title) or "Documento")
